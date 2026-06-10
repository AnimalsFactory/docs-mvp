from flask import Flask, render_template, request, redirect, url_for, jsonify, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document as DocxDocument
import io
from bs4 import BeautifulSoup

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///documents.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = 'supersecretkey123'

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# ====================== МОДЕЛИ ======================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), default="Без названия")
    content = db.Column(db.Text, default="")
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

with app.app_context():
    db.create_all()

# ====================== МАРШРУТЫ ======================
@app.route('/')
@login_required
def index():
    docs = Document.query.filter_by(user_id=current_user.id).all()
    return render_template('index.html', documents=docs)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if User.query.filter_by(username=username).first():
            flash('Пользователь уже существует')
            return redirect(url_for('register'))
        user = User(username=username, password_hash=generate_password_hash(password))
        db.session.add(user)
        db.session.commit()
        flash('Регистрация прошла успешно!')
        return redirect(url_for('login'))
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for('index'))
        flash('Неверный логин или пароль')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/create', methods=['POST'])
@login_required
def create():
    doc = Document(title="Без названия", content="", user_id=current_user.id)
    db.session.add(doc)
    db.session.commit()
    return redirect(url_for('editor', doc_id=doc.id))

@app.route('/editor/<int:doc_id>')
@login_required
def editor(doc_id):
    doc = Document.query.get_or_404(doc_id)
    if doc.user_id != current_user.id:
        return redirect(url_for('index'))
    return render_template('editor.html', document=doc)

@app.route('/save/<int:doc_id>', methods=['POST'])
@login_required
def save(doc_id):
    doc = Document.query.get_or_404(doc_id)
    if doc.user_id != current_user.id:
        return jsonify({"status": "error"}), 403
    doc.title = request.form.get('title', 'Без названия')
    doc.content = request.form.get('content', '')
    db.session.commit()
    return jsonify({"status": "success"})

@app.route('/delete/<int:doc_id>', methods=['POST'])
@login_required
def delete(doc_id):
    doc = Document.query.get_or_404(doc_id)
    if doc.user_id == current_user.id:
        db.session.delete(doc)
        db.session.commit()
    return redirect(url_for('index'))

# ====================== УЛУЧШЕННЫЙ ЭКСПОРТ В DOCX ======================
@app.route('/export/docx/<int:doc_id>')
@login_required
def export_docx(doc_id):
    doc = Document.query.get_or_404(doc_id)
    if doc.user_id != current_user.id:
        return redirect(url_for('index'))
    
    word_doc = DocxDocument()
    word_doc.add_heading(doc.title, 0)
    
    # Очистка HTML от Quill
    soup = BeautifulSoup(doc.content, 'html.parser')
    
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'strong', 'em', 'u']):
        text = element.get_text().strip()
        if not text:
            continue
            
        if element.name.startswith('h'):
            level = int(element.name[1])
            word_doc.add_heading(text, level)
        elif element.name == 'li':
            word_doc.add_paragraph(text, style='List Bullet')
        else:
            para = word_doc.add_paragraph(text)
            # Простое форматирование
            if element.find('strong') or element.find('b'):
                for run in para.runs:
                    run.bold = True
            if element.find('em') or element.find('i'):
                for run in para.runs:
                    run.italic = True
    
    # Добавляем изображения (если они есть)
    for img in soup.find_all('img'):
        word_doc.add_paragraph("[Изображение вставлено]")
    
    file_stream = io.BytesIO()
    word_doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{doc.title}.docx"
    )

if __name__ == '__main__':
    app.run(debug=True)